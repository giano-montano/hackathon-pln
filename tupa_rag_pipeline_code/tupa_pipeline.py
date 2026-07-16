#!/usr/bin/env python3
"""
Pipeline de preprocesamiento del TUPA SUNAT 2018 para un sistema RAG.

Etapas:
1. Convierte el .doc antiguo a .docx y el índice .xls a .xlsx.
2. Lee la relación oficial de procedimientos para recuperar jerarquía y nombres canónicos.
3. Extrae la tabla consolidada del .docx preservando las 14 columnas del TUPA.
4. Reconstruye procedimientos y subprocedimientos, incluidos los servicios prestados
   en exclusividad.
5. Genera documentos padre estructurados en JSONL.
6. Opcionalmente genera chunks semánticos listos para indexación RAG.
7. Emite un reporte de calidad y archivos de diagnóstico.

El pipeline es extractivo: no inventa ni completa requisitos, plazos, costos o autoridades.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import logging
import os
import platform
import re
import shutil
import subprocess
import sys
import tempfile
import unicodedata
import zipfile
from collections import Counter, OrderedDict, defaultdict
from dataclasses import asdict, dataclass, field
from pathlib import Path
from statistics import mean, median
from typing import Any, Iterable, Iterator, Mapping, Sequence

from docx import Document
from openpyxl import load_workbook


SCHEMA_VERSION = "1.0.0"
LOGGER = logging.getLogger("tupa_pipeline")

# Esquema real de la tabla consolidada convertida a DOCX.
# El documento tiene 14 columnas de cuadrícula.
TUPA_2018_COLUMNS: tuple[str, ...] = (
    "codigo",
    "denominacion_y_base_legal",
    "requisitos",
    "formularios",
    "costo",
    "calificacion_automatica",
    "evaluacion_previa_positiva",
    "evaluacion_previa_negativa",
    "plazo",
    "inicio_procedimiento",
    "autoridad_competente",
    "reconsideracion",
    "reclamo",
    "apelacion",
)

CONTENT_FIELDS: tuple[str, ...] = (
    "fundamento_legal",
    "requisitos",
    "formularios",
    "costo",
    "calificacion",
    "plazo",
    "inicio_procedimiento",
    "autoridad_competente",
    "reconsideracion",
    "reclamo",
    "apelacion",
)

FIELD_LABELS: Mapping[str, str] = {
    "fundamento_legal": "Fundamento legal",
    "requisitos": "Requisitos",
    "formularios": "Formularios, códigos y ubicaciones",
    "costo": "Derecho de tramitación / costo",
    "calificacion": "Calificación",
    "plazo": "Plazo para resolver",
    "inicio_procedimiento": "Inicio del procedimiento / canal",
    "autoridad_competente": "Autoridad competente",
    "reconsideracion": "Reconsideración",
    "reclamo": "Reclamo",
    "apelacion": "Apelación",
    "denominacion_adicional": "Información adicional de la denominación",
}

KNOWN_SEMANTIC_HEADINGS = (
    "REQUISITOS GENERALES",
    "REQUISITOS ESPECÍFICOS",
    "DOCUMENTACIÓN - REQUISITOS GENERALES",
    "DOCUMENTACIÓN - REQUISITOS ESPECÍFICOS",
    "DOCUMENTACION - REQUISITOS GENERALES",
    "DOCUMENTACION - REQUISITOS ESPECIFICOS",
    "CONDICIONES PARA PRESENTAR",
    "CONDICIONES PARA OBTENER",
    "CONDICIONES PARA SER HABILITADO",
    "EN FORMA PRESENCIAL",
    "A TRAVÉS DE SUNAT VIRTUAL",
    "A TRAVES DE SUNAT VIRTUAL",
    "EN SUNAT VIRTUAL",
    "DE MANERA PRESENCIAL",
    "A TRAVÉS DEL CANAL TELEFÓNICO",
    "A TRAVES DEL CANAL TELEFONICO",
    "REQUISITOS",
    "NOTA",
    "NOTAS",
)

STATUS_KEYWORDS = (
    "ELIMINADO MEDIANTE",
    "MODIFICADO MEDIANTE",
    "MODIFICADO A SERVICIO",
    "INCORPORADO MEDIANTE",
)

SUBPROCEDURE_RE = re.compile(
    r"^\s*0*(?P<code>\d+(?:-[A-Z])?(?:\.\d+)+)\s*(?:[.\-–—:]+\s*)?(?P<title>.*)$",
    re.IGNORECASE | re.DOTALL,
)

CODE_RE = re.compile(r"^(?:0*(\d+)(?:\s*-\s*([A-Z]))?|SERVICIO\s*0*(\d+))$", re.IGNORECASE)


class PipelineError(RuntimeError):
    """Error controlado del pipeline."""


@dataclass
class IndexSubprocedure:
    code: str
    title: str
    raw_title: str
    source_row: int


@dataclass
class IndexEntry:
    code: str
    title: str
    raw_title: str
    section: str | None
    category: str | None
    status: str | None
    source_row: int
    subprocedures: list[IndexSubprocedure] = field(default_factory=list)


@dataclass
class TableRow:
    table_index: int
    row_index: int
    cells: list[str]


@dataclass
class ProcedureAccumulator:
    code: str
    rows: list[TableRow] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Normalización y utilidades
# ---------------------------------------------------------------------------


def clean_text(value: Any) -> str:
    """Normaliza texto sin alterar su contenido semántico."""
    if value is None:
        return ""
    text = str(value)
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\xa0", " ").replace("\x07", " ")
    text = text.replace("\t", " ")
    # Elimina caracteres de control, conservando saltos de línea.
    text = "".join(ch if ch == "\n" or ord(ch) >= 32 else " " for ch in text)
    lines: list[str] = []
    for line in text.split("\n"):
        line = re.sub(r"[ ]+", " ", line).strip()
        lines.append(line)
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def strip_accents(text: str) -> str:
    decomposed = unicodedata.normalize("NFD", text)
    return "".join(ch for ch in decomposed if unicodedata.category(ch) != "Mn")


def normalize_for_match(value: Any) -> str:
    text = strip_accents(clean_text(value)).upper()
    text = re.sub(r"[^A-Z0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def normalize_code(value: Any) -> str:
    """Normaliza códigos: 01 -> 1, 08-A -> 8-A, Servicio 01 -> SERVICIO 1."""
    text = clean_text(value).upper().strip(" .|\n")
    text = text.replace("N.°", "").replace("N°", "").replace("Nº", "")
    text = re.sub(r"\s+", " ", text)
    # Excel puede devolver números como 1.0.
    if re.fullmatch(r"\d+\.0+", text):
        text = text.split(".", 1)[0]
    match = CODE_RE.fullmatch(text)
    if not match:
        return text
    numeric, suffix, service = match.groups()
    if service is not None:
        return f"SERVICIO {int(service)}"
    assert numeric is not None
    code = str(int(numeric))
    return f"{code}-{suffix.upper()}" if suffix else code


def normalize_subprocedure_code(value: str) -> str:
    value = clean_text(value).upper().strip()
    parts = value.split(".")
    if not parts:
        return value
    parent = normalize_code(parts[0])
    rest = [str(int(p)) if p.isdigit() else p for p in parts[1:]]
    return ".".join([parent, *rest])


def slugify(value: str, max_length: int = 80) -> str:
    text = strip_accents(clean_text(value)).lower()
    text = re.sub(r"[^a-z0-9]+", "_", text).strip("_")
    if not text:
        text = "item"
    return text[:max_length].rstrip("_")


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as file_obj:
        for block in iter(lambda: file_obj.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def fingerprint(text: str) -> str:
    return hashlib.sha1(normalize_for_match(text).encode("utf-8")).hexdigest()


def unique_blocks(values: Iterable[str]) -> list[str]:
    """Conserva bloques no vacíos en orden, eliminando repeticiones exactas normalizadas."""
    result: list[str] = []
    seen: set[str] = set()
    for value in values:
        value = clean_text(value)
        if not value:
            continue
        key = fingerprint(value)
        if key in seen:
            continue
        seen.add(key)
        result.append(value)
    return result


def join_unique_blocks(values: Iterable[str]) -> str:
    return "\n\n".join(unique_blocks(values)).strip()


def detect_status_and_clean_title(raw_title: str) -> tuple[str, str | None]:
    """Separa el nombre del procedimiento de notas como ELIMINADO/MODIFICADO."""
    raw_title = clean_text(raw_title)
    lines = raw_title.splitlines()
    status_start: int | None = None
    for index, line in enumerate(lines):
        normalized = normalize_for_match(line)
        if any(keyword in normalized for keyword in STATUS_KEYWORDS):
            status_start = index
            break
    if status_start is not None:
        title = " ".join(lines[:status_start]).strip()
        status = " ".join(lines[status_start:]).strip()
        title = title.strip("() ")
        return title or raw_title, status

    # Algunas notas aparecen entre paréntesis en la misma línea.
    match = re.search(
        r"\s*\((ELIMINADO|MODIFICADO|INCORPORADO)\s+MEDIANTE.+$",
        raw_title,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if match:
        return raw_title[: match.start()].strip(), raw_title[match.start() :].strip(" ()")
    return raw_title, None


def detect_subprocedure(raw_text: str) -> tuple[str, str] | None:
    match = SUBPROCEDURE_RE.match(clean_text(raw_text))
    if not match:
        return None
    code = normalize_subprocedure_code(match.group("code"))
    title = clean_text(match.group("title"))
    return code, title


def approximate_token_count(text: str) -> int:
    """Estimador local; no depende del tokenizer del modelo."""
    return len(re.findall(r"\w+|[^\w\s]", text, flags=re.UNICODE))


def compress_ranges(numbers: Iterable[int]) -> list[str]:
    ordered = sorted(set(numbers))
    if not ordered:
        return []
    ranges: list[str] = []
    start = previous = ordered[0]
    for number in ordered[1:]:
        if number == previous + 1:
            previous = number
            continue
        ranges.append(str(start) if start == previous else f"{start}-{previous}")
        start = previous = number
    ranges.append(str(start) if start == previous else f"{start}-{previous}")
    return ranges


# ---------------------------------------------------------------------------
# Conversión de archivos antiguos
# ---------------------------------------------------------------------------


def find_soffice(explicit_path: str | None = None) -> str | None:
    if explicit_path:
        candidate = Path(explicit_path)
        if candidate.exists():
            return str(candidate)
        located = shutil.which(explicit_path)
        if located:
            return located
    for name in ("soffice", "libreoffice"):
        located = shutil.which(name)
        if located:
            return located
    return None


def run_libreoffice_conversion(
    source: Path,
    output_dir: Path,
    target_extension: str,
    soffice_path: str,
    timeout_seconds: int,
) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    target_extension = target_extension.lower().lstrip(".")
    format_spec = {
        "docx": "docx:Office Open XML Text",
        "xlsx": "xlsx:Calc MS Excel 2007 XML",
        "html": "html:HTML (StarWriter)",
    }.get(target_extension, target_extension)

    expected = output_dir / f"{source.stem}.{target_extension}"
    if expected.exists():
        expected.unlink()

    with tempfile.TemporaryDirectory(prefix="tupa_lo_profile_") as profile_dir:
        profile_uri = Path(profile_dir).resolve().as_uri()
        command = [
            soffice_path,
            f"-env:UserInstallation={profile_uri}",
            "--headless",
            "--norestore",
            "--nodefault",
            "--nolockcheck",
            "--convert-to",
            format_spec,
            "--outdir",
            str(output_dir.resolve()),
            str(source.resolve()),
        ]
        LOGGER.info("Convirtiendo %s a .%s con LibreOffice", source.name, target_extension)
        try:
            completed = subprocess.run(
                command,
                check=False,
                capture_output=True,
                text=True,
                timeout=timeout_seconds,
            )
        except subprocess.TimeoutExpired as exc:
            raise PipelineError(
                f"La conversión de {source.name} excedió {timeout_seconds} segundos. "
                "Aumenta --conversion-timeout; este TUPA puede tardar varios minutos."
            ) from exc

    if completed.returncode != 0:
        raise PipelineError(
            "LibreOffice no pudo convertir el archivo.\n"
            f"Comando: {' '.join(command)}\n"
            f"Salida: {completed.stdout}\nError: {completed.stderr}"
        )

    # LibreOffice puede variar ligeramente el nombre; buscar el archivo recién generado.
    if not expected.exists():
        candidates = sorted(output_dir.glob(f"{source.stem}*.{target_extension}"))
        if len(candidates) == 1:
            expected = candidates[0]
        else:
            raise PipelineError(
                f"La conversión terminó, pero no se encontró el archivo .{target_extension}. "
                f"Salida de LibreOffice: {completed.stdout} {completed.stderr}"
            )
    if expected.stat().st_size == 0:
        raise PipelineError(f"El archivo convertido está vacío: {expected}")
    return expected


def convert_doc_with_word_com(source: Path, destination: Path) -> Path:
    """Fallback para Windows con Microsoft Word y pywin32 instalados."""
    try:
        import pythoncom  # type: ignore
        import win32com.client  # type: ignore
    except ImportError as exc:
        raise PipelineError(
            "No se encontró LibreOffice y tampoco pywin32 para usar Microsoft Word."
        ) from exc

    destination.parent.mkdir(parents=True, exist_ok=True)
    pythoncom.CoInitialize()
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(str(source.resolve()), ReadOnly=True)
        # 16 = wdFormatDocumentDefault (.docx)
        document.SaveAs2(str(destination.resolve()), FileFormat=16)
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()
    if not destination.exists():
        raise PipelineError("Microsoft Word no generó el archivo DOCX esperado.")
    return destination


def convert_xls_with_excel_com(source: Path, destination: Path) -> Path:
    """Fallback para Windows con Microsoft Excel y pywin32 instalados."""
    try:
        import pythoncom  # type: ignore
        import win32com.client  # type: ignore
    except ImportError as exc:
        raise PipelineError(
            "No se encontró LibreOffice y tampoco pywin32 para usar Microsoft Excel."
        ) from exc

    destination.parent.mkdir(parents=True, exist_ok=True)
    pythoncom.CoInitialize()
    excel = None
    workbook = None
    try:
        excel = win32com.client.DispatchEx("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        workbook = excel.Workbooks.Open(str(source.resolve()), ReadOnly=True)
        # 51 = xlOpenXMLWorkbook (.xlsx)
        workbook.SaveAs(str(destination.resolve()), FileFormat=51)
    finally:
        if workbook is not None:
            workbook.Close(False)
        if excel is not None:
            excel.Quit()
        pythoncom.CoUninitialize()
    if not destination.exists():
        raise PipelineError("Microsoft Excel no generó el archivo XLSX esperado.")
    return destination


def ensure_docx(
    source: Path,
    converted_dir: Path,
    soffice_path: str | None,
    timeout_seconds: int,
    reuse: bool = True,
) -> Path:
    source = source.resolve()
    if source.suffix.lower() == ".docx":
        return source
    if source.suffix.lower() != ".doc":
        raise PipelineError(f"Se esperaba un archivo .doc o .docx: {source}")

    destination = converted_dir / f"{source.stem}.docx"
    if reuse and destination.exists() and destination.stat().st_mtime >= source.stat().st_mtime:
        LOGGER.info("Reutilizando DOCX convertido: %s", destination)
        return destination

    if soffice_path:
        destination = run_libreoffice_conversion(
            source, converted_dir, "docx", soffice_path, timeout_seconds
        )
    elif platform.system() == "Windows":
        LOGGER.info("LibreOffice no está disponible; intentando Microsoft Word.")
        destination = convert_doc_with_word_com(source, destination)
    else:
        raise PipelineError(
            "No se encontró LibreOffice (soffice). Instálalo o proporciona --soffice."
        )

    if not zipfile.is_zipfile(destination):
        raise PipelineError(f"El DOCX convertido no es un ZIP/Office Open XML válido: {destination}")
    return destination


def ensure_xlsx(
    source: Path,
    converted_dir: Path,
    soffice_path: str | None,
    timeout_seconds: int,
    reuse: bool = True,
) -> Path:
    source = source.resolve()
    if source.suffix.lower() == ".xlsx":
        return source
    if source.suffix.lower() != ".xls":
        raise PipelineError(f"Se esperaba un archivo .xls o .xlsx: {source}")

    destination = converted_dir / f"{source.stem}.xlsx"
    if reuse and destination.exists() and destination.stat().st_mtime >= source.stat().st_mtime:
        LOGGER.info("Reutilizando XLSX convertido: %s", destination)
        return destination

    if soffice_path:
        destination = run_libreoffice_conversion(
            source, converted_dir, "xlsx", soffice_path, timeout_seconds
        )
    elif platform.system() == "Windows":
        LOGGER.info("LibreOffice no está disponible; intentando Microsoft Excel.")
        destination = convert_xls_with_excel_com(source, destination)
    else:
        raise PipelineError(
            "No se encontró LibreOffice (soffice). Instálalo o proporciona --soffice."
        )
    if not zipfile.is_zipfile(destination):
        raise PipelineError(f"El XLSX convertido no es válido: {destination}")
    return destination


# ---------------------------------------------------------------------------
# Índice oficial de procedimientos
# ---------------------------------------------------------------------------


def is_index_header(code_value: Any, title: str) -> bool:
    code_norm = normalize_for_match(code_value)
    title_norm = normalize_for_match(title)
    return (
        "NUMERO DEL PROCEDIMIENTO TUPA" in code_norm
        or "NOMBRE DEL PROCEDIMIENTO ADMINISTRATIVO" in title_norm
    )


def parse_index_workbook(path: Path) -> OrderedDict[str, IndexEntry]:
    workbook = load_workbook(path, data_only=True, read_only=True)
    worksheet = workbook.active

    entries: OrderedDict[str, IndexEntry] = OrderedDict()
    current_section: str | None = None
    current_category: str | None = None
    current_entry: IndexEntry | None = None

    for row_number, row in enumerate(worksheet.iter_rows(values_only=True), start=1):
        code_value = row[0] if len(row) > 0 else None
        title_value = row[1] if len(row) > 1 else None
        title = clean_text(title_value)
        if not title:
            continue
        if is_index_header(code_value, title):
            continue

        title_match = normalize_for_match(title)
        if title_match.startswith("SECCION "):
            current_section = title
            current_category = None
            current_entry = None
            continue
        if title_match.startswith("NOV 2018"):
            continue

        code = normalize_code(code_value) if code_value is not None else ""
        if code:
            clean_title, status = detect_status_and_clean_title(title)
            entry = IndexEntry(
                code=code,
                title=clean_title,
                raw_title=title,
                section=current_section,
                category=current_category,
                status=status,
                source_row=row_number,
            )
            if code in entries:
                raise PipelineError(
                    f"Código duplicado en el índice: {code}, filas "
                    f"{entries[code].source_row} y {row_number}"
                )
            entries[code] = entry
            current_entry = entry
            continue

        detected = detect_subprocedure(title)
        if detected and current_entry is not None:
            sub_code, sub_title = detected
            parent_code = normalize_code(sub_code.split(".", 1)[0])
            if parent_code == normalize_code(current_entry.code):
                current_entry.subprocedures.append(
                    IndexSubprocedure(
                        code=sub_code,
                        title=sub_title,
                        raw_title=title,
                        source_row=row_number,
                    )
                )
                continue

        # Una fila sin código que no es subprocedimiento es un encabezado temático.
        current_category = title
        current_entry = None

    if not entries:
        raise PipelineError(f"No se encontraron procedimientos en el índice: {path}")
    LOGGER.info(
        "Índice cargado: %d entradas y %d subprocedimientos",
        len(entries),
        sum(len(entry.subprocedures) for entry in entries.values()),
    )
    return entries


# ---------------------------------------------------------------------------
# Extracción y reconstrucción del TUPA
# ---------------------------------------------------------------------------


def extract_docx_table_rows(path: Path) -> tuple[list[TableRow], dict[str, Any]]:
    document = Document(path)
    rows: list[TableRow] = []
    table_shapes: list[dict[str, int]] = []

    for table_index, table in enumerate(document.tables):
        row_lengths = [len(row.cells) for row in table.rows]
        if not row_lengths:
            continue
        min_columns = min(row_lengths)
        max_columns = max(row_lengths)
        table_shapes.append(
            {
                "table_index": table_index,
                "rows": len(table.rows),
                "min_columns": min_columns,
                "max_columns": max_columns,
            }
        )
        if min_columns != len(TUPA_2018_COLUMNS) or max_columns != len(TUPA_2018_COLUMNS):
            raise PipelineError(
                f"La tabla {table_index} tiene entre {min_columns} y {max_columns} columnas; "
                f"se esperaban {len(TUPA_2018_COLUMNS)}. Revisa la conversión del DOC."
            )
        for row_index, row in enumerate(table.rows):
            cells = [clean_text(cell.text) for cell in row.cells]
            rows.append(TableRow(table_index=table_index, row_index=row_index, cells=cells))

    if not rows:
        raise PipelineError(f"No se encontraron tablas en {path}")
    metadata = {
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "table_shapes": table_shapes,
        "row_count": len(rows),
    }
    LOGGER.info("DOCX extraído: %d tablas y %d filas", len(document.tables), len(rows))
    return rows, metadata


def row_text(row: TableRow) -> str:
    return "\n".join(cell for cell in row.cells if cell)


def is_service_header_row(row: TableRow) -> bool:
    normalized = normalize_for_match(" ".join(row.cells[:2]))
    return "SUNAT" in normalized and "SERVICIOS PRESTADOS EN EXCLUSIVIDAD" in normalized


def is_tupa_header_row(row: TableRow) -> bool:
    normalized = normalize_for_match(" ".join(row.cells[:2]))
    return "SUNAT" in normalized and "TEXTO UNICO DE PROCEDIMIENTOS ADMINISTRATIVOS" in normalized


def is_column_header_row(row: TableRow) -> bool:
    normalized = normalize_for_match(row_text(row))
    keywords = (
        "NUMERO DE ORDEN",
        "DENOMINACION DEL PROCEDIMIENTO",
        "DERECHO DE TRAMITACION",
        "AUTORIDAD COMPETENTE",
    )
    return sum(keyword in normalized for keyword in keywords) >= 2


def match_service_code_by_title(
    row: TableRow, index_entries: Mapping[str, IndexEntry]
) -> str | None:
    denomination = row.cells[1] if len(row.cells) > 1 else ""
    denomination_norm = normalize_for_match(denomination)
    if not denomination_norm:
        return None
    for code, entry in index_entries.items():
        if not code.startswith("SERVICIO "):
            continue
        title_norm = normalize_for_match(entry.title)
        if title_norm and (denomination_norm.startswith(title_norm) or title_norm.startswith(denomination_norm[:80])):
            return code
    return None


def group_rows_by_procedure(
    rows: Sequence[TableRow], index_entries: Mapping[str, IndexEntry]
) -> tuple[OrderedDict[str, ProcedureAccumulator], dict[str, Any]]:
    known_codes = set(index_entries)
    groups: OrderedDict[str, ProcedureAccumulator] = OrderedDict()
    current_code: str | None = None
    mode = "procedure"
    unknown_rows: list[dict[str, Any]] = []
    orphan_rows: list[dict[str, Any]] = []
    header_rows: list[dict[str, int]] = []

    for row in rows:
        if is_service_header_row(row):
            mode = "service"
            current_code = None
            header_rows.append({"table": row.table_index, "row": row.row_index})
            continue
        if is_tupa_header_row(row):
            mode = "procedure"
            current_code = None
            header_rows.append({"table": row.table_index, "row": row.row_index})
            continue
        if is_column_header_row(row):
            current_code = None
            header_rows.append({"table": row.table_index, "row": row.row_index})
            continue
        if not any(row.cells):
            continue

        raw_code = normalize_code(row.cells[0])
        code: str | None = None
        if mode == "service":
            if raw_code.isdigit() and 1 <= int(raw_code) <= 4:
                candidate = f"SERVICIO {int(raw_code)}"
                if candidate in known_codes:
                    code = candidate
            if code is None:
                code = match_service_code_by_title(row, index_entries)
        elif raw_code in known_codes:
            code = raw_code

        if code is not None:
            current_code = code
            accumulator = groups.setdefault(code, ProcedureAccumulator(code=code))
            accumulator.rows.append(row)
            continue

        if raw_code:
            unknown_rows.append(
                {
                    "table": row.table_index,
                    "row": row.row_index,
                    "raw_code": raw_code,
                    "denomination_preview": clean_text(row.cells[1])[:250],
                }
            )
            current_code = None
            continue

        # Las filas sin código suelen ser continuaciones o subprocedimientos del procedimiento anterior.
        if current_code is not None:
            groups[current_code].rows.append(row)
        else:
            orphan_rows.append(
                {
                    "table": row.table_index,
                    "row": row.row_index,
                    "preview": row_text(row)[:250],
                }
            )

    diagnostics = {
        "unknown_rows": unknown_rows,
        "orphan_rows": orphan_rows,
        "header_rows": header_rows,
    }
    return groups, diagnostics


def flexible_title_prefix_pattern(title: str) -> re.Pattern[str] | None:
    words = re.findall(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ0-9$%/°º.\-]+", clean_text(title))
    if not words:
        return None
    escaped = [re.escape(word) for word in words]
    return re.compile(r"^\s*" + r"[\s\n]+".join(escaped), re.IGNORECASE)


def split_denomination_and_legal(
    denomination_text: str, canonical_title: str
) -> tuple[str, str, str]:
    """
    Retorna (titulo_detectado, fundamento_legal, denominacion_adicional).
    La denominación canónica siempre proviene del índice; esta función solo separa el contenido.
    """
    text = clean_text(denomination_text)
    if not text:
        return "", "", ""

    title_detected = ""
    remainder = text
    prefix = flexible_title_prefix_pattern(canonical_title)
    if prefix:
        match = prefix.match(text)
        if match:
            title_detected = clean_text(text[: match.end()])
            remainder = clean_text(text[match.end() :])

    marker = re.search(
        r"FUNDAMENTOS?\s+LEGAL(?:ES)?\s*:?",
        remainder,
        flags=re.IGNORECASE,
    )
    if marker:
        before = clean_text(remainder[: marker.start()])
        legal = clean_text(remainder[marker.end() :])
        return title_detected, legal, before

    # Si se eliminó el título canónico, el resto suele ser directamente la base legal.
    legal_start = re.search(
        r"(?mi)^(DECRETO|RESOLUCI[ÓO]N|LEY\b|DECISI[ÓO]N|CONVENIO|ACUERDO|REGLAMENTO)",
        remainder,
    )
    if legal_start:
        additional = clean_text(remainder[: legal_start.start()])
        legal = clean_text(remainder[legal_start.start() :])
        return title_detected, legal, additional

    if title_detected:
        return title_detected, "", remainder

    # Fallback con marcador localizado en el texto completo.
    marker_full = re.search(
        r"FUNDAMENTOS?\s+LEGAL(?:ES)?\s*:?",
        text,
        flags=re.IGNORECASE,
    )
    if marker_full:
        return (
            clean_text(text[: marker_full.start()]),
            clean_text(text[marker_full.end() :]),
            "",
        )
    return "", "", text


def has_x(value: str) -> bool:
    normalized = normalize_for_match(value)
    return bool(re.search(r"(?:^|\s)X(?:$|\s)", normalized))


def derive_calification(field_values: Mapping[str, str]) -> tuple[str, dict[str, bool]]:
    flags = {
        "automatico": has_x(field_values.get("calificacion_automatica", "")),
        "evaluacion_previa_positiva": has_x(
            field_values.get("evaluacion_previa_positiva", "")
        ),
        "evaluacion_previa_negativa": has_x(
            field_values.get("evaluacion_previa_negativa", "")
        ),
    }
    labels: list[str] = []
    if flags["automatico"]:
        labels.append("Automático")
    if flags["evaluacion_previa_positiva"]:
        labels.append("Evaluación previa con silencio positivo")
    if flags["evaluacion_previa_negativa"]:
        labels.append("Evaluación previa con silencio negativo")

    # Si no hay X, conservar los valores originales sin interpretar.
    if not labels:
        originals = unique_blocks(
            [
                field_values.get("calificacion_automatica", ""),
                field_values.get("evaluacion_previa_positiva", ""),
                field_values.get("evaluacion_previa_negativa", ""),
            ]
        )
        labels.extend(originals)
    return "; ".join(labels), flags


def detect_document_status(text: str) -> str | None:
    normalized = normalize_for_match(text)
    if "ELIMINADO MEDIANTE" in normalized:
        for line in clean_text(text).splitlines():
            if "ELIMINADO" in normalize_for_match(line):
                return line
        return "Eliminado según la fuente"
    return None


def build_parent_page_content(record: Mapping[str, Any]) -> str:
    metadata = record["metadata"]
    fields = record["fields"]
    entity_label = "Servicio" if str(record["code"]).startswith("SERVICIO ") else "Procedimiento TUPA"
    lines = [
        f"# {entity_label} {record['code']}: {record['title']}",
    ]
    if metadata.get("section"):
        lines.append(f"Sección: {metadata['section']}")
    if metadata.get("category"):
        lines.append(f"Categoría: {metadata['category']}")
    if metadata.get("status"):
        lines.append(f"Estado en la fuente: {metadata['status']}")
    subprocedures = metadata.get("subprocedures") or []
    if subprocedures:
        lines.append("Subprocedimientos del índice:")
        for item in subprocedures:
            lines.append(f"- {item['code']}: {item['title']}")

    for field_name in CONTENT_FIELDS:
        value = clean_text(fields.get(field_name, ""))
        if not value:
            continue
        lines.append(f"\n## {FIELD_LABELS[field_name]}\n{value}")
    additional = clean_text(fields.get("denominacion_adicional", ""))
    if additional:
        lines.append(f"\n## {FIELD_LABELS['denominacion_adicional']}\n{additional}")
    return "\n".join(lines).strip()


def build_procedure_records(
    groups: Mapping[str, ProcedureAccumulator],
    index_entries: Mapping[str, IndexEntry],
    source_doc: Path,
    converted_docx: Path,
) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []

    # Salida en el orden canónico del índice, no en el orden accidental del documento.
    for code, index_entry in index_entries.items():
        accumulator = groups.get(code)
        if accumulator is None:
            continue

        values_by_column: dict[str, list[str]] = defaultdict(list)
        for row in accumulator.rows:
            for column_index, field_name in enumerate(TUPA_2018_COLUMNS):
                if field_name == "codigo":
                    continue
                value = row.cells[column_index]
                if value:
                    values_by_column[field_name].append(value)

        joined = {
            field_name: join_unique_blocks(values_by_column.get(field_name, []))
            for field_name in TUPA_2018_COLUMNS
            if field_name != "codigo"
        }
        title_detected, legal, denomination_additional = split_denomination_and_legal(
            joined.get("denominacion_y_base_legal", ""), index_entry.title
        )
        calification, calification_flags = derive_calification(joined)

        combined_text = join_unique_blocks(
            [joined.get(name, "") for name in joined if joined.get(name)]
        )
        document_status = detect_document_status(combined_text)
        status = index_entry.status or document_status

        source_rows_by_table: dict[int, list[int]] = defaultdict(list)
        for row in accumulator.rows:
            source_rows_by_table[row.table_index].append(row.row_index)
        source_trace = [
            {
                "table_index": table_index,
                "row_ranges_zero_based": compress_ranges(row_numbers),
                "row_ranges_human": compress_ranges(number + 1 for number in row_numbers),
            }
            for table_index, row_numbers in sorted(source_rows_by_table.items())
        ]

        fields = {
            "fundamento_legal": legal,
            "requisitos": joined.get("requisitos", ""),
            "formularios": joined.get("formularios", ""),
            "costo": joined.get("costo", ""),
            "calificacion": calification,
            "plazo": joined.get("plazo", ""),
            "inicio_procedimiento": joined.get("inicio_procedimiento", ""),
            "autoridad_competente": joined.get("autoridad_competente", ""),
            "reconsideracion": joined.get("reconsideracion", ""),
            "reclamo": joined.get("reclamo", ""),
            "apelacion": joined.get("apelacion", ""),
            "denominacion_adicional": denomination_additional,
        }
        subprocedures = [asdict(item) for item in index_entry.subprocedures]
        record: dict[str, Any] = {
            "schema_version": SCHEMA_VERSION,
            "id": f"tupa_{slugify(code)}",
            "code": code,
            "title": index_entry.title,
            "metadata": {
                "document_type": "servicio_prestado_en_exclusividad"
                if code.startswith("SERVICIO ")
                else "procedimiento_tupa",
                "section": index_entry.section,
                "category": index_entry.category,
                "status": status,
                "index_source_row": index_entry.source_row,
                "subprocedures": subprocedures,
                "calification_flags": calification_flags,
                "source_file": source_doc.name,
                "converted_file": converted_docx.name,
                "source_trace": source_trace,
            },
            "fields": fields,
            "extraction": {
                "title_detected_in_document": title_detected,
                "row_count": len(accumulator.rows),
            },
        }
        record["page_content"] = build_parent_page_content(record)
        records.append(record)
    return records


# ---------------------------------------------------------------------------
# Chunking semántico
# ---------------------------------------------------------------------------


def official_subprocedure_map(
    subprocedures: Sequence[Mapping[str, Any]],
) -> dict[str, tuple[str, str]]:
    result: dict[str, tuple[str, str]] = {}
    for sub in subprocedures:
        raw_code = clean_text(sub.get("code", ""))
        if not raw_code:
            continue
        code = normalize_subprocedure_code(raw_code)
        result[code] = (raw_code, clean_text(sub.get("title", "")))
    return result


def line_official_subprocedure(
    line: str, subprocedures: Sequence[Mapping[str, Any]]
) -> tuple[str, str] | None:
    detected = detect_subprocedure(line)
    if not detected:
        return None
    detected_code, _ = detected
    official = official_subprocedure_map(subprocedures)
    return official.get(detected_code)


def is_semantic_heading(
    line: str, subprocedures: Sequence[Mapping[str, Any]]
) -> bool:
    normalized = normalize_for_match(line)
    if not normalized:
        return False
    # Solo los subprocedimientos presentes en el índice oficial se consideran jerarquía.
    # Así se evita confundir fechas (18.09.2004) o numerales internos (5.1, 7.2)
    # con subprocedimientos del TUPA.
    if line_official_subprocedure(line, subprocedures):
        return True
    for heading in KNOWN_SEMANTIC_HEADINGS:
        heading_norm = normalize_for_match(heading)
        if normalized == heading_norm or normalized.startswith(heading_norm + " "):
            return len(line.split()) <= 35
    return False


def semantic_sections(
    text: str, subprocedures: Sequence[Mapping[str, Any]]
) -> list[tuple[str | None, str]]:
    """Divide por subprocedimientos oficiales y encabezados explícitos."""
    text = clean_text(text)
    if not text:
        return []

    sections: list[tuple[str | None, str]] = []
    current_heading: str | None = None
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_lines
        body = clean_text("\n".join(current_lines))
        if body or current_heading:
            full = body
            if current_heading and body:
                full = f"{current_heading}\n{body}"
            elif current_heading:
                full = current_heading
            sections.append((current_heading, clean_text(full)))
        current_lines = []

    for line in text.splitlines():
        stripped = clean_text(line)
        if stripped and is_semantic_heading(stripped, subprocedures):
            flush()
            current_heading = stripped
        else:
            current_lines.append(stripped)
    flush()
    return [(heading, body) for heading, body in sections if body]


def split_oversized_unit(text: str, max_tokens: int) -> list[str]:
    text = clean_text(text)
    if approximate_token_count(text) <= max_tokens:
        return [text]

    sentences = [
        clean_text(piece)
        for piece in re.split(r"(?<=[.;!?])\s+(?=[A-ZÁÉÍÓÚÜÑ0-9])", text)
        if clean_text(piece)
    ]
    if len(sentences) <= 1:
        words = text.split()
        chunks: list[str] = []
        current: list[str] = []
        for word in words:
            tentative = " ".join([*current, word])
            if current and approximate_token_count(tentative) > max_tokens:
                chunks.append(" ".join(current))
                current = [word]
            else:
                current.append(word)
        if current:
            chunks.append(" ".join(current))
        return chunks

    chunks = []
    current_sentences: list[str] = []
    for sentence in sentences:
        tentative = " ".join([*current_sentences, sentence])
        if current_sentences and approximate_token_count(tentative) > max_tokens:
            chunks.append(" ".join(current_sentences))
            current_sentences = [sentence]
        else:
            current_sentences.append(sentence)
    if current_sentences:
        chunks.append(" ".join(current_sentences))
    # Una oración extraordinariamente larga todavía puede superar el máximo.
    final: list[str] = []
    for chunk in chunks:
        if approximate_token_count(chunk) > max_tokens:
            final.extend(split_oversized_unit(chunk, max_tokens))
        else:
            final.append(chunk)
    return final


def pack_units(
    units: Sequence[str], target_tokens: int, max_tokens: int
) -> list[str]:
    expanded: list[str] = []
    for unit in units:
        expanded.extend(split_oversized_unit(unit, max_tokens))

    packed: list[str] = []
    current: list[str] = []
    for unit in expanded:
        tentative = "\n\n".join([*current, unit])
        tentative_count = approximate_token_count(tentative)
        if current and tentative_count > max_tokens:
            packed.append("\n\n".join(current))
            current = [unit]
            continue
        current.append(unit)
        if approximate_token_count("\n\n".join(current)) >= target_tokens:
            packed.append("\n\n".join(current))
            current = []
    if current:
        packed.append("\n\n".join(current))
    return [clean_text(item) for item in packed if clean_text(item)]


def identify_subprocedure(
    heading_or_text: str | None, subprocedures: Sequence[Mapping[str, Any]]
) -> tuple[str | None, str | None]:
    if not heading_or_text:
        return None, None
    official = line_official_subprocedure(heading_or_text, subprocedures)
    if official:
        return official

    normalized = normalize_for_match(heading_or_text)
    for sub in subprocedures:
        title = clean_text(sub.get("title", ""))
        if title and normalize_for_match(title) in normalized:
            return clean_text(sub.get("code", "")), title
    return None, None


def chunk_prefix(
    record: Mapping[str, Any],
    section_label: str,
    subprocedure_code: str | None,
    subprocedure_title: str | None,
) -> str:
    metadata = record["metadata"]
    entity_label = "Servicio" if str(record["code"]).startswith("SERVICIO ") else "Procedimiento TUPA"
    lines = [
        "SUNAT - TUPA 2018",
        f"{entity_label}: {record['code']}",
        f"Nombre: {record['title']}",
    ]
    if metadata.get("section"):
        lines.append(f"Sección: {metadata['section']}")
    if metadata.get("category"):
        lines.append(f"Categoría: {metadata['category']}")
    if metadata.get("status"):
        lines.append(f"Estado en la fuente: {metadata['status']}")
    if subprocedure_code:
        sub_line = f"Subprocedimiento: {subprocedure_code}"
        if subprocedure_title:
            sub_line += f" - {subprocedure_title}"
        lines.append(sub_line)
    lines.append(f"Tipo de contenido: {section_label}")
    return "\n".join(lines)


def make_chunk(
    record: Mapping[str, Any],
    section_name: str,
    section_label: str,
    content: str,
    sequence: int,
    subprocedure_code: str | None = None,
    subprocedure_title: str | None = None,
) -> dict[str, Any]:
    prefix = chunk_prefix(
        record, section_label, subprocedure_code, subprocedure_title
    )
    page_content = f"{prefix}\n\n{clean_text(content)}".strip()
    return {
        "schema_version": SCHEMA_VERSION,
        "id": f"{record['id']}_{slugify(section_name)}_{sequence:03d}",
        "parent_id": record["id"],
        "metadata": {
            "codigo_tupa": record["code"],
            "procedimiento": record["title"],
            "document_type": record["metadata"].get("document_type"),
            "section": record["metadata"].get("section"),
            "category": record["metadata"].get("category"),
            "status": record["metadata"].get("status"),
            "content_type": section_name,
            "content_label": section_label,
            "subprocedure_code": subprocedure_code,
            "subprocedure_title": subprocedure_title,
            "source_file": record["metadata"].get("source_file"),
            "chunk_sequence": sequence,
            "approx_tokens": approximate_token_count(page_content),
        },
        "page_content": page_content,
        "raw_text": clean_text(content),
    }


def build_chunks_for_record(
    record: Mapping[str, Any], target_tokens: int, max_tokens: int
) -> list[dict[str, Any]]:
    if target_tokens <= 0 or max_tokens <= target_tokens:
        raise PipelineError("Debe cumplirse 0 < target_tokens < max_tokens.")

    chunks: list[dict[str, Any]] = []
    sequence_by_section: Counter[str] = Counter()
    subprocedures = record["metadata"].get("subprocedures") or []

    long_fields = ("fundamento_legal", "requisitos", "formularios")
    for field_name in long_fields:
        value = clean_text(record["fields"].get(field_name, ""))
        if not value:
            continue
        sections = semantic_sections(value, subprocedures) or [(None, value)]
        for heading, body in sections:
            sub_code, sub_title = identify_subprocedure(heading or body[:250], subprocedures)
            # Los bloques semánticos se empacan por párrafos; nunca se corta por caracteres.
            paragraphs = [
                clean_text(part)
                for part in re.split(r"\n\s*\n", body)
                if clean_text(part)
            ]
            if not paragraphs:
                paragraphs = [body]
            prefix_tokens = approximate_token_count(
                chunk_prefix(record, FIELD_LABELS[field_name], sub_code, sub_title)
            )
            content_max = max(40, max_tokens - prefix_tokens - 4)
            content_target = max(30, min(content_max, target_tokens - prefix_tokens - 4))
            for packed in pack_units(paragraphs, content_target, content_max):
                sequence_by_section[field_name] += 1
                chunks.append(
                    make_chunk(
                        record,
                        field_name,
                        FIELD_LABELS[field_name],
                        packed,
                        sequence_by_section[field_name],
                        sub_code,
                        sub_title,
                    )
                )

    grouped_fields = (
        (
            "resumen_administrativo",
            "Costo, calificación y plazo",
            ("costo", "calificacion", "plazo"),
        ),
        (
            "canal_y_autoridad",
            "Canal de inicio y autoridad competente",
            ("inicio_procedimiento", "autoridad_competente"),
        ),
        (
            "recursos",
            "Recursos administrativos",
            ("reconsideracion", "reclamo", "apelacion"),
        ),
        (
            "informacion_adicional",
            "Información adicional",
            ("denominacion_adicional",),
        ),
    )

    for section_name, section_label, fields in grouped_fields:
        blocks: list[str] = []
        for field_name in fields:
            value = clean_text(record["fields"].get(field_name, ""))
            if value:
                blocks.append(f"{FIELD_LABELS.get(field_name, field_name)}:\n{value}")
        if not blocks:
            continue
        content = "\n\n".join(blocks)
        # Estos bloques suelen ser cortos, pero se protegen ante casos muy extensos.
        prefix_tokens = approximate_token_count(
            chunk_prefix(record, section_label, None, None)
        )
        content_max = max(40, max_tokens - prefix_tokens - 4)
        content_target = max(30, min(content_max, target_tokens - prefix_tokens - 4))
        for packed in pack_units([content], content_target, content_max):
            sequence_by_section[section_name] += 1
            chunks.append(
                make_chunk(
                    record,
                    section_name,
                    section_label,
                    packed,
                    sequence_by_section[section_name],
                )
            )

    if not chunks:
        sequence_by_section["documento_completo"] += 1
        chunks.append(
            make_chunk(
                record,
                "documento_completo",
                "Contenido completo",
                record["page_content"],
                sequence_by_section["documento_completo"],
            )
        )
    return chunks


def build_all_chunks(
    records: Sequence[Mapping[str, Any]], target_tokens: int, max_tokens: int
) -> list[dict[str, Any]]:
    chunks: list[dict[str, Any]] = []
    for record in records:
        chunks.extend(build_chunks_for_record(record, target_tokens, max_tokens))
    return chunks


# ---------------------------------------------------------------------------
# Validación y escritura
# ---------------------------------------------------------------------------


def write_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    with temporary.open("w", encoding="utf-8") as file_obj:
        json.dump(value, file_obj, ensure_ascii=False, indent=2)
        file_obj.write("\n")
    temporary.replace(path)


def write_jsonl(path: Path, records: Iterable[Mapping[str, Any]]) -> int:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    count = 0
    with temporary.open("w", encoding="utf-8") as file_obj:
        for record in records:
            file_obj.write(json.dumps(record, ensure_ascii=False) + "\n")
            count += 1
    temporary.replace(path)
    return count


def write_csv(path: Path, rows: Sequence[Mapping[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if not rows:
        path.write_text("", encoding="utf-8")
        return
    fieldnames: list[str] = []
    seen: set[str] = set()
    for row in rows:
        for key in row:
            if key not in seen:
                fieldnames.append(key)
                seen.add(key)
    with path.open("w", encoding="utf-8", newline="") as file_obj:
        writer = csv.DictWriter(file_obj, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def index_catalog(entries: Mapping[str, IndexEntry]) -> list[dict[str, Any]]:
    return [asdict(entry) for entry in entries.values()]


def build_quality_report(
    source_doc: Path,
    source_index: Path,
    converted_docx: Path,
    converted_xlsx: Path,
    index_entries: Mapping[str, IndexEntry],
    groups: Mapping[str, ProcedureAccumulator],
    records: Sequence[Mapping[str, Any]],
    chunks: Sequence[Mapping[str, Any]],
    extraction_metadata: Mapping[str, Any],
    grouping_diagnostics: Mapping[str, Any],
) -> dict[str, Any]:
    expected_codes = list(index_entries)
    extracted_codes = [str(record["code"]) for record in records]
    missing_codes = [code for code in expected_codes if code not in groups]
    extra_codes = [code for code in groups if code not in index_entries]

    empty_fields: dict[str, list[str]] = defaultdict(list)
    for record in records:
        for field_name in CONTENT_FIELDS:
            if not clean_text(record["fields"].get(field_name, "")):
                empty_fields[field_name].append(str(record["code"]))

    chunk_tokens = [
        int(chunk["metadata"].get("approx_tokens", 0)) for chunk in chunks
    ]
    duplicate_chunk_hashes: dict[str, list[str]] = defaultdict(list)
    for chunk in chunks:
        duplicate_chunk_hashes[fingerprint(chunk["page_content"])].append(chunk["id"])
    duplicates = [ids for ids in duplicate_chunk_hashes.values() if len(ids) > 1]

    warnings: list[str] = []
    if missing_codes:
        warnings.append(f"Faltan {len(missing_codes)} códigos del índice en el TUPA extraído.")
    if grouping_diagnostics.get("unknown_rows"):
        warnings.append(
            f"Hay {len(grouping_diagnostics['unknown_rows'])} filas con código no reconocido."
        )
    if grouping_diagnostics.get("orphan_rows"):
        warnings.append(
            f"Hay {len(grouping_diagnostics['orphan_rows'])} filas huérfanas sin procedimiento padre."
        )
    if duplicates:
        warnings.append(f"Se detectaron {len(duplicates)} grupos de chunks duplicados.")

    report: dict[str, Any] = {
        "schema_version": SCHEMA_VERSION,
        "input_files": {
            "tupa_source": {
                "path": str(source_doc),
                "sha256": sha256_file(source_doc),
                "bytes": source_doc.stat().st_size,
            },
            "index_source": {
                "path": str(source_index),
                "sha256": sha256_file(source_index),
                "bytes": source_index.stat().st_size,
            },
            "converted_docx": {
                "path": str(converted_docx),
                "sha256": sha256_file(converted_docx),
                "bytes": converted_docx.stat().st_size,
            },
            "converted_xlsx": {
                "path": str(converted_xlsx),
                "sha256": sha256_file(converted_xlsx),
                "bytes": converted_xlsx.stat().st_size,
            },
        },
        "extraction": dict(extraction_metadata),
        "coverage": {
            "index_entries": len(index_entries),
            "grouped_entries": len(groups),
            "parent_records": len(records),
            "coverage_ratio": round(len(records) / len(index_entries), 6)
            if index_entries
            else 0,
            "missing_codes": missing_codes,
            "extra_codes": extra_codes,
            "unknown_rows": grouping_diagnostics.get("unknown_rows", []),
            "orphan_rows": grouping_diagnostics.get("orphan_rows", []),
            "header_rows": grouping_diagnostics.get("header_rows", []),
        },
        "content": {
            "empty_field_counts": {
                key: len(value) for key, value in sorted(empty_fields.items())
            },
            "empty_field_codes": dict(empty_fields),
            "status_counts": dict(
                Counter(
                    clean_text(record["metadata"].get("status")) or "vigente_en_documento"
                    for record in records
                )
            ),
            "row_count_per_procedure": {
                record["code"]: record["extraction"]["row_count"] for record in records
            },
        },
        "chunks": {
            "count": len(chunks),
            "approx_tokens_min": min(chunk_tokens) if chunk_tokens else 0,
            "approx_tokens_max": max(chunk_tokens) if chunk_tokens else 0,
            "approx_tokens_mean": round(mean(chunk_tokens), 2) if chunk_tokens else 0,
            "approx_tokens_median": median(chunk_tokens) if chunk_tokens else 0,
            "over_650_tokens": sum(value > 650 for value in chunk_tokens),
            "duplicate_groups": duplicates,
        },
        "warnings": warnings,
        "result": "ok" if not missing_codes and not grouping_diagnostics.get("unknown_rows") else "review",
    }
    return report


def write_preview(path: Path, records: Sequence[Mapping[str, Any]], limit: int = 5) -> None:
    lines = [
        "# Vista previa del corpus TUPA procesado",
        "",
        f"Se muestran los primeros {min(limit, len(records))} documentos padre.",
        "",
    ]
    for record in records[:limit]:
        lines.append(record["page_content"])
        lines.append("\n---\n")
    path.write_text("\n".join(lines), encoding="utf-8")


def write_debug_rows(path: Path, rows: Sequence[TableRow]) -> None:
    debug_records = []
    for row in rows:
        item: dict[str, Any] = {
            "table_index": row.table_index,
            "row_index": row.row_index,
        }
        for index, field_name in enumerate(TUPA_2018_COLUMNS):
            item[field_name] = row.cells[index]
        debug_records.append(item)
    write_jsonl(path, debug_records)


# ---------------------------------------------------------------------------
# Orquestación
# ---------------------------------------------------------------------------


def run_pipeline(args: argparse.Namespace) -> dict[str, Path]:
    source_doc = Path(args.tupa_doc).expanduser().resolve()
    source_index = Path(args.index_xls).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve()

    if not source_doc.exists():
        raise PipelineError(f"No existe el TUPA: {source_doc}")
    if not source_index.exists():
        raise PipelineError(f"No existe el índice: {source_index}")

    if args.force and output_dir.exists():
        shutil.rmtree(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    converted_dir = output_dir / "converted"
    converted_dir.mkdir(parents=True, exist_ok=True)

    soffice_path = find_soffice(args.soffice)
    if soffice_path:
        LOGGER.info("LibreOffice encontrado: %s", soffice_path)

    converted_docx = ensure_docx(
        source_doc,
        converted_dir,
        soffice_path,
        args.conversion_timeout,
        reuse=not args.no_reuse_converted,
    )
    converted_xlsx = ensure_xlsx(
        source_index,
        converted_dir,
        soffice_path,
        args.conversion_timeout,
        reuse=not args.no_reuse_converted,
    )

    entries = parse_index_workbook(converted_xlsx)
    rows, extraction_metadata = extract_docx_table_rows(converted_docx)
    groups, grouping_diagnostics = group_rows_by_procedure(rows, entries)
    records = build_procedure_records(
        groups, entries, source_doc=source_doc, converted_docx=converted_docx
    )

    chunks: list[dict[str, Any]] = []
    if not args.no_chunks:
        chunks = build_all_chunks(records, args.target_tokens, args.max_tokens)

    paths = {
        "index_catalog": output_dir / "index_catalog.json",
        "parents": output_dir / "procedures.jsonl",
        "chunks": output_dir / "chunks.jsonl",
        "quality_report": output_dir / "quality_report.json",
        "preview": output_dir / "preview.md",
        "diagnostics_unknown": output_dir / "diagnostics_unknown_rows.csv",
        "diagnostics_orphan": output_dir / "diagnostics_orphan_rows.csv",
    }

    write_json(paths["index_catalog"], index_catalog(entries))
    write_jsonl(paths["parents"], records)
    if not args.no_chunks:
        write_jsonl(paths["chunks"], chunks)
    elif paths["chunks"].exists():
        paths["chunks"].unlink()
    write_preview(paths["preview"], records, args.preview_count)
    write_csv(paths["diagnostics_unknown"], grouping_diagnostics["unknown_rows"])
    write_csv(paths["diagnostics_orphan"], grouping_diagnostics["orphan_rows"])

    if args.debug_rows:
        debug_path = output_dir / "debug_extracted_rows.jsonl"
        write_debug_rows(debug_path, rows)
        paths["debug_rows"] = debug_path

    report = build_quality_report(
        source_doc,
        source_index,
        converted_docx,
        converted_xlsx,
        entries,
        groups,
        records,
        chunks,
        extraction_metadata,
        grouping_diagnostics,
    )
    report["configuration"] = {
        "target_tokens": args.target_tokens,
        "max_tokens": args.max_tokens,
        "chunks_enabled": not args.no_chunks,
        "conversion_timeout_seconds": args.conversion_timeout,
    }
    report["output_files"] = {key: str(value) for key, value in paths.items()}
    write_json(paths["quality_report"], report)

    LOGGER.info(
        "Pipeline completado: %d procedimientos padre y %d chunks. Cobertura %.1f%%",
        len(records),
        len(chunks),
        100 * report["coverage"]["coverage_ratio"],
    )
    return paths


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Preprocesa el TUPA SUNAT 2018 para un sistema RAG.",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    parser.add_argument("--tupa-doc", required=True, help="Ruta a tupa_consolidado.doc o .docx")
    parser.add_argument(
        "--index-xls", required=True, help="Ruta a relacionTupa-2018.xls o .xlsx"
    )
    parser.add_argument("--output-dir", default="output_tupa", help="Directorio de salida")
    parser.add_argument(
        "--soffice",
        default=None,
        help="Ruta al ejecutable soffice/libreoffice; se detecta automáticamente",
    )
    parser.add_argument(
        "--conversion-timeout",
        type=int,
        default=1200,
        help="Tiempo máximo por conversión. El DOC consolidado puede tardar varios minutos.",
    )
    parser.add_argument(
        "--target-tokens",
        type=int,
        default=450,
        help="Tamaño objetivo aproximado de cada chunk incluyendo su contexto",
    )
    parser.add_argument(
        "--max-tokens",
        type=int,
        default=650,
        help="Tamaño máximo aproximado de cada chunk incluyendo su contexto",
    )
    parser.add_argument("--no-chunks", action="store_true", help="Solo genera documentos padre")
    parser.add_argument(
        "--no-reuse-converted",
        action="store_true",
        help="Fuerza la reconversión aunque ya exista un archivo convertido reciente",
    )
    parser.add_argument(
        "--debug-rows",
        action="store_true",
        help="Guarda cada fila extraída para depuración",
    )
    parser.add_argument(
        "--preview-count", type=int, default=5, help="Número de procedimientos en preview.md"
    )
    parser.add_argument(
        "--force", action="store_true", help="Elimina el directorio de salida antes de ejecutar"
    )
    parser.add_argument(
        "--log-level",
        choices=("DEBUG", "INFO", "WARNING", "ERROR"),
        default="INFO",
    )
    return parser


def main() -> int:
    parser = build_arg_parser()
    args = parser.parse_args()
    logging.basicConfig(
        level=getattr(logging, args.log_level),
        format="%(asctime)s | %(levelname)s | %(message)s",
    )
    try:
        paths = run_pipeline(args)
    except (PipelineError, OSError, ValueError) as exc:
        LOGGER.error("Pipeline detenido: %s", exc)
        return 1
    print("\nArchivos generados:")
    for name, path in paths.items():
        if path.exists():
            print(f"- {name}: {path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
